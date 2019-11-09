from docx import Document
from docx.shared import Pt, Cm
import os
from openpyxl import load_workbook
from tkinter import *

# initial page
window = Tk()
window.title("Sheet Generator 0.1")
window.geometry('300x80')
spreadsheet_filename = Label(window, text="Spreadsheet filename:")
spreadsheet_filename.grid(column=0, row=0)
entry_spreadsheet_filename = Entry(window)
entry_spreadsheet_filename.insert(0, 'spreadsheet.xlsx')
entry_spreadsheet_filename.grid(column=1, row=0)
sheet_name = Label(window, text="Sheet:")
sheet_name.grid(column=0, row=1)
entry_sheet_name = Entry(window)
entry_sheet_name.insert(0, 'Sheet1')
entry_sheet_name.grid(column=1, row=1)


def generate_forms():
    parent_file = entry_spreadsheet_filename.get()
    wb = load_workbook(filename=parent_file)
    ws1 = wb[entry_sheet_name.get()]

    # Define worker fields
    class Worker:
        def __init__(self, first_name, family_name, gender, birthday, citizenship, worker_role, identity, contact,
                     picture_file):
            self.first_name = first_name
            self.family_name = family_name
            self.gender = gender
            self.birthday = birthday
            self.citizenship = citizenship
            self.worker_role = worker_role
            self.identity = identity
            self.contact = contact
            self.picture_file = picture_file

    # Check how many workers inside spreadsheet
    def count_workers(sheet):
        flag = True
        number_of_workers = 0
        while flag is True:
            if sheet['B' + str(number_of_workers + 1)].value is None:
                flag = False
            else:
                number_of_workers += 1
        return number_of_workers

    # function to avoid None
    def xstr(s):
        return '' if s is None else str(s)

    # Create a list of class objects
    workers = list()
    for i in range(count_workers(ws1) - 1):
        workers.append(Worker(
            # first name
            ws1['C' + str(i + 2)].value,
            # family name
            ws1['B' + str(i + 2)].value,
            # gender
            ws1['F' + str(i + 2)].value,
            # birthday
            str(ws1['G' + str(i + 2)].value),
            # citizenship
            ws1['H' + str(i + 2)].value,
            # worker role
            str(ws1['E' + str(i + 2)].value),
            # identity
            {'type': xstr(ws1['I' + str(i + 2)].value),
             'number': xstr(ws1['K' + str(i + 2)].value),
             'issuing agency': xstr(ws1['J' + str(i + 2)].value),
             'issue date': xstr(ws1['L' + str(i + 2)].value),
             'expiry date': xstr(ws1['M' + str(i + 2)].value)},
            # contact
            xstr(ws1['Q' + str(i + 2)].value),
            # Picture
            xstr(ws1['AB' + str(i + 2)].value),
        ))

    # Create sheet based on template forms
    document = Document(os.getcwd() + '\\forms.docx')
    table = document.tables[0]
    header = table.cell(0, 0)

    # Adding standard style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(8)

    # citizenship
    def add_citizenship(nat):
        citizenship = table.cell(0, 4)
        citizenship.paragraphs[3].text = ' ' * 10 + nat

    # job position
    def add_worker_role(role):
        worker_role = table.cell(0, 2)
        worker_role.paragraphs[4].text += ' ' * 70 + role

    # family name
    def add_family_name(name):
        family_name = table.cell(0, 6)
        family_name.paragraphs[3].text += ' ' * 15 + name

    # given first name
    def add_first_name(name):
        first_name = table.cell(0, 8)
        first_name.paragraphs[3].text += ' ' * 15 + name

    # gender
    def add_gender(selected_gender):
        gender = table.cell(0, 10)
        if selected_gender == 'M' or selected_gender == 'm':
            gender.tables[1].cell(0, 2).text = 'X'
        elif selected_gender == 'F' or selected_gender == 'f':
            gender.tables[1].cell(0, 0).text = 'X'
        else:
            gender.tables[1].cell(0, 4).text = selected_gender

    # birthday //should be string
    def add_birthday(day):
        birthday = table.cell(0, 10)
        birthday.tables[1].cell(0, 26).text = day[3]
        birthday.tables[1].cell(0, 25).text = day[2]
        birthday.tables[1].cell(0, 24).text = day[1]
        birthday.tables[1].cell(0, 23).text = day[0]
        birthday.tables[1].cell(0, 21).text = day[6]
        birthday.tables[1].cell(0, 20).text = day[5]
        birthday.tables[1].cell(0, 18).text = day[9]
        birthday.tables[1].cell(0, 17).text = day[8]

    # Passport // should be a dictionary

    def add_identity(identity):
        identity_number = table.cell(0, 12)
        identity_number.paragraphs[4].text = 'type: ' + identity['type'] + ' ' * 5 + 'number: ' + identity[
            'number'] + ' ' * 5 + 'issuing agency: ' + identity['issuing agency'] + ' ' * 50 + 'issue date: ' + \
                                             identity['issue date'] + ' ' * 5 + 'expiry date: ' + identity[
                                                 'expiry date']

    # Contact //should be string // 16 empty boxes
    def add_contact(phone_number):
        contact = table.cell(0, 14)
        contact.tables[0].cell(0, 0).text = '+'
        contact.tables[0].cell(0, 1).text = '6'
        contact.tables[0].cell(0, 2).text = '3'
        contact.tables[0].cell(0, 3).text = ' '
        for x in range(len(phone_number)):
            contact.tables[0].cell(0, x + 3).text = phone_number[x]

    def add_picture(picture_filename):
        picture = table.cell(0, 1)
        picture.paragraphs[1].text += ' ' * 5
        picture.paragraphs[1].add_run().add_picture(picture_filename, width=Cm(3.5))

    if not os.path.exists(os.getcwd() + '\\output'):
        os.mkdir(os.getcwd() + '\\output')

    for person in workers:
        # first_name, family_name, gender, birthday, citizenship, worker_role, identity, contact,picture_file
        document = Document(os.getcwd() + '\\forms.docx')
        table = document.tables[0]
        add_first_name(person.first_name)
        add_family_name(person.family_name)
        add_gender(person.gender)
        add_birthday(person.birthday)
        add_citizenship(person.citizenship)
        add_worker_role(person.worker_role)
        add_identity(person.identity)
        add_contact(person.contact)
        add_picture(os.getcwd() + '\\pictures\\' + person.picture_file)
        if os.path.exists(
                os.getcwd() + '\\output\\' + person.first_name + '_' + person.family_name + '_ceremonies.docx'):
            os.remove(os.getcwd() + '\\output\\' + person.first_name + '_' + person.family_name + '_ceremonies.docx')
        document.save(os.getcwd() + '\\output\\' + person.first_name + '_' + person.family_name + '_ceremonies.docx')

    window.destroy()


btn = Button(window, text="Generate", command=generate_forms)
btn.grid(column=1, row=2)
window.mainloop()
